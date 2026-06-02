"""
Credo AI Branded Document Builder
Generates a .docx file with the official Credo AI brand design system:
  - Logo embedded in header (every page, top-right)
  - Decorative swirls on odd pages only (top-left + bottom-right, behind text)
  - Plus Jakarta Sans / Instrument Sans typography
  - Brand colors: #7B04CA (purple), #161616 (text), #F2DEFF (callout bg)
  - Page number footer

Usage (from Bash via python3):
    python3 build_doc.py <config_json_path>

The config JSON should match the schema in the SKILL.md examples.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import os, glob, sys, json

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'assets')
LOGO_PATH = os.path.join(ASSETS_DIR, 'credo-logo.png')
SWIRL_PATH = os.path.join(ASSETS_DIR, 'swirl.png')
SWIRL_FLIP = os.path.join(ASSETS_DIR, 'swirl-flipped.png')

# Page dimensions in EMU (US Letter)
PAGE_W = 7772400
PAGE_H = 10058400
SWIRL_W = 1280160   # ~1.4"
SWIRL_H = 5852160   # ~6.4"

# ─── Style helpers ────────────────────────────────────────────────────

def hex_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def color_run(run, hex_color):
    r, g, b = hex_rgb(hex_color)
    run.font.color.rgb = RGBColor(r, g, b)

def set_para_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.lstrip('#'))
    pPr.append(shd)

def add_hrule(para, color='D6D6D6'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)

def find_font():
    for name in ['Instrument Sans', 'Plus Jakarta Sans']:
        key = name.replace(' ', '')
        for d in [os.path.expanduser('~/Library/Fonts'),
                  '/Library/Fonts',
                  '/System/Library/Fonts']:
            if glob.glob(os.path.join(d, f'*{key}*')):
                return name
    return 'Calibri'

def styled_run(para, text, font, size, bold=False, color='161616', italic=False):
    r = para.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    color_run(r, f'#{color}')
    return r

# ─── Odd/even header setup ────────────────────────────────────────────

def enable_even_odd_headers(doc):
    """Enable separate headers for even vs odd pages."""
    settings = doc.settings.element
    existing = settings.find(qn('w:evenAndOddHeaders'))
    if existing is not None:
        settings.remove(existing)
    eo = OxmlElement('w:evenAndOddHeaders')
    settings.append(eo)
    for s in doc.sections:
        s.different_first_page_header_footer = False
        sectPr = s._sectPr
        t = sectPr.find(qn('w:titlePg'))
        if t is not None:
            sectPr.remove(t)

# ─── Anchored image (positioned absolutely on page, behind text) ──────

_drawing_id = [100]
def next_id():
    _drawing_id[0] += 1
    return _drawing_id[0]

def add_anchored_image(paragraph, image_path, width_emu, height_emu, x_offset_emu, y_offset_emu):
    run = paragraph.add_run()
    run.add_picture(image_path, width=Emu(width_emu))
    drawing = run._r.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))
    graphic = inline.find(qn('a:graphic'))
    rid = next_id()
    anchor_xml = f'''<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
                xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                distT="0" distB="0" distL="0" distR="0"
                simplePos="0" relativeHeight="{rid}" behindDoc="1" locked="0"
                layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>{x_offset_emu}</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>{y_offset_emu}</wp:posOffset></wp:positionV>
        <wp:extent cx="{width_emu}" cy="{height_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="{rid}" name="Decoration{rid}"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
    </wp:anchor>'''
    anchor = parse_xml(anchor_xml)
    anchor.append(graphic)
    drawing.remove(inline)
    drawing.append(anchor)

# ─── Header / Footer ──────────────────────────────────────────────────

def populate_header(header, font, logo_path, with_swirls=False):
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.clear()
    if with_swirls:
        add_anchored_image(para, SWIRL_PATH, SWIRL_W, SWIRL_H,
                           x_offset_emu=-457200, y_offset_emu=-457200)
        add_anchored_image(para, SWIRL_FLIP, SWIRL_W, SWIRL_H,
                           x_offset_emu=PAGE_W - SWIRL_W + 457200,
                           y_offset_emu=PAGE_H - SWIRL_H + 457200)
    logo_run = para.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.8))

def populate_footer(footer, font):
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.clear()
    run = para.add_run()
    fld = OxmlElement('w:fldChar'); fld.set(qn('w:fldCharType'), 'begin'); run._r.append(fld)
    r2 = para.add_run()
    it = OxmlElement('w:instrText'); it.text = 'PAGE'; r2._r.append(it)
    r3 = para.add_run()
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end'); r3._r.append(fe)
    for r in [run, r2, r3]:
        r.font.name = font; r.font.size = Pt(9)
        color_run(r, '#9A9A9A')

def setup_headers_footers(doc, font, swirls_on_odd_pages=True):
    section = doc.sections[0]
    section.header_distance = Inches(0.4)
    # Odd page header (default) — swirls if enabled
    odd_header = section.header
    odd_header.is_linked_to_previous = False
    populate_header(odd_header, font, LOGO_PATH, with_swirls=swirls_on_odd_pages)
    # Even page header — never swirls
    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    populate_header(even_header, font, LOGO_PATH, with_swirls=False)
    # Footers — both same
    populate_footer(section.footer, font)
    populate_footer(section.even_page_footer, font)

# ─── Content block renderers ──────────────────────────────────────────

def render_block(doc, block, font):
    btype = block.get('type', 'body')
    text = block.get('text', '')

    if btype == 'signpost':
        p = doc.add_paragraph()
        styled_run(p, text, font, 9, color='8A8A8A')
        p.paragraph_format.space_after = Pt(4)
        return

    if btype == 'title':
        p = doc.add_paragraph()
        styled_run(p, text, font, 36, bold=True)
        p.paragraph_format.space_after = Pt(6)
        return

    if btype == 'subtitle':
        p = doc.add_paragraph()
        styled_run(p, text, font, 13, color='8A8A8A')
        p.paragraph_format.space_after = Pt(10)
        return

    if btype == 'rule':
        p = doc.add_paragraph()
        add_hrule(p, block.get('color', 'D6D6D6'))
        p.paragraph_format.space_before = Pt(block.get('space_before', 8))
        p.paragraph_format.space_after = Pt(block.get('space_after', 12))
        return

    if btype in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        style_map = {
            'h1': (22, True,  '161616'),
            'h2': (16, True,  '161616'),
            'h3': (13, True,  '161616'),
            'h4': (11, True,  '161616'),
            'h5': (10, False, '8A8A8A'),
            'h6': (9,  False, 'AAAAAA'),
        }
        size, bold, color = style_map[btype]
        p = doc.add_paragraph()
        styled_run(p, text, font, size, bold=bold, color=color)
        if btype in ('h1', 'h2'):
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        return

    if btype == 'callout':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        set_para_shading(p, 'F2DEFF')
        styled_run(p, text, font, 11)
        return

    if btype == 'bullet':
        p = doc.add_paragraph(style='List Bullet')
        styled_run(p, text, font, 11)
        return

    if btype == 'bullet_bold':
        # Format: "Bold lead, rest of text" — split at first comma+space
        p = doc.add_paragraph(style='List Bullet')
        lead = block.get('lead', '')
        rest = block.get('rest', '')
        if lead:
            r1 = p.add_run(lead); r1.font.name = font; r1.font.size = Pt(11); r1.font.bold = True
            color_run(r1, '#161616')
        if rest:
            r2 = p.add_run(' ' + rest); r2.font.name = font; r2.font.size = Pt(11)
            color_run(r2, '#161616')
        return

    if btype == 'lead_para':
        # Bold lead-in followed by regular body
        p = doc.add_paragraph()
        styled_run(p, block.get('lead', ''), font, 11, bold=True)
        styled_run(p, block.get('rest', ''), font, 11)
        p.paragraph_format.space_after = Pt(10)
        return

    if btype == 'stat':
        # Big purple number + label
        p = doc.add_paragraph()
        r1 = p.add_run(block.get('number', '') + '  '); r1.font.name = font; r1.font.size = Pt(14); r1.font.bold = True
        color_run(r1, '#7B04CA')
        r2 = p.add_run(block.get('label', '')); r2.font.name = font; r2.font.size = Pt(11)
        color_run(r2, '#161616')
        p.paragraph_format.space_after = Pt(4)
        return

    if btype == 'quote':
        # Indented italic quote in muted gray
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        styled_run(p, text, font, 11, italic=True, color='5A5A5A')
        p.paragraph_format.space_after = Pt(10)
        return

    if btype == 'caption':
        p = doc.add_paragraph()
        styled_run(p, text, font, 9, italic=True, color='9A9A9A')
        p.paragraph_format.space_after = Pt(12)
        return

    if btype == 'attribution':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, text, font, 9, color='9A9A9A')
        return

    # Default: body paragraph (supports inline bold/italic via 'parts' array)
    p = doc.add_paragraph()
    parts = block.get('parts')
    if parts:
        for part in parts:
            styled_run(p, part.get('text', ''), font, 11,
                       bold=part.get('bold', False),
                       italic=part.get('italic', False),
                       color=part.get('color', '161616'))
    else:
        styled_run(p, text, font, 11)
    p.paragraph_format.space_after = Pt(block.get('space_after', 10))

# ─── Main ─────────────────────────────────────────────────────────────

def build(config):
    font = find_font()
    print(f"Font: {font}")
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.25)
        s.right_margin = Inches(1.25)
    enable_even_odd_headers(doc)
    setup_headers_footers(doc, font, swirls_on_odd_pages=config.get('swirls', True))

    for block in config.get('blocks', []):
        render_block(doc, block, font)

    out = os.path.expanduser(config.get('output_path',
                                        f"~/Desktop/{config.get('title', 'Credo_AI_Document').replace(' ', '_')}.docx"))
    doc.save(out)
    print(f"Saved: {out}")
    return out

if __name__ == '__main__':
    if len(sys.argv) > 1:
        with open(sys.argv[1]) as f:
            config = json.load(f)
    else:
        print("Usage: python3 build_doc.py <config_json_path>", file=sys.stderr)
        sys.exit(1)
    build(config)
